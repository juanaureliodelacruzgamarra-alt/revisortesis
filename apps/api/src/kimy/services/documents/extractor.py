"""Document text extraction.

Word: uses python-docx, preserves heading levels via paragraph style.
PDF:  uses pypdf, returns plain text per page (no style info available).
"""
from __future__ import annotations

import io
from dataclasses import dataclass

from docx import Document as DocxDocument
from pypdf import PdfReader


@dataclass(slots=True)
class ExtractedParagraph:
    text: str
    heading_level: int | None  # 1..9 for Word headings, None for body / PDF lines


@dataclass(slots=True)
class ExtractedDocument:
    paragraphs: list[ExtractedParagraph]
    page_count: int  # 0 if unknown (.docx)

    @property
    def full_text(self) -> str:
        return "\n".join(p.text for p in self.paragraphs if p.text)


def extract_from_docx(content: bytes) -> ExtractedDocument:
    doc = DocxDocument(io.BytesIO(content))
    out: list[ExtractedParagraph] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = (p.style.name or "").strip()
        level: int | None = None
        if style.lower().startswith("heading"):
            try:
                level = int(style.split()[-1])
            except (ValueError, IndexError):
                level = 1
        elif style.lower() in {"title", "subtitle"}:
            level = 1
        out.append(ExtractedParagraph(text=text, heading_level=level))
    return ExtractedDocument(paragraphs=out, page_count=0)


def extract_from_pdf(content: bytes) -> ExtractedDocument:
    reader = PdfReader(io.BytesIO(content))
    out: list[ExtractedParagraph] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        for raw in text.splitlines():
            line = raw.strip()
            if line:
                out.append(ExtractedParagraph(text=line, heading_level=None))
    return ExtractedDocument(paragraphs=out, page_count=len(reader.pages))


def extract(filename: str, content: bytes, mime_type: str) -> ExtractedDocument:
    name = filename.lower()
    if name.endswith(".docx") or "wordprocessingml" in mime_type:
        return extract_from_docx(content)
    if name.endswith(".pdf") or mime_type == "application/pdf":
        return extract_from_pdf(content)
    raise ValueError(f"Unsupported file type: {filename!r} ({mime_type!r})")
